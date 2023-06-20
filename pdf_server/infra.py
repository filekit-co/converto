import io
import logging
from typing import Optional

import fitz
from docx import Document
from pdf2docx import Converter
from pdf2docx.page.Pages import Pages

from utils import out_filename


class DocxConverter(Converter):
    """pdf2docx converter adding with statement."""
    
    def __init__(self, file_name:str, input_file_type:str, input_file:bytes, path_or_stream: bytes, password:str=None):
        '''Initialize fitz object with given pdf file path.

        Args:
            pdf_file (str): pdf file path.
            password (str): Password for encrypted pdf. Default to None if not encrypted.
        '''
        # fitz object
        self.filename_pdf = file_name
        self.password = str(password or '')
        self._fitz_doc = fitz.Document(stream=input_file, filetype=input_file_type)
        self.path_or_stream = path_or_stream

        # initialize empty pages container
        self._pages = Pages()


    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

    @property
    def docx_name(self):
        return out_filename(self.filename_pdf, '.pdf', '.docx')
        

    @property
    def doc_name(self):
        return out_filename(self.filename_pdf, '.pdf', '.doc')

    def convert(self, docx_filename:str=None, start:int=0, end:int=None, pages:list=None, **kwargs):
        """Convert specified PDF pages to docx file.

        Args:
            docx_filename (str, optional): docx filename to write to. Defaults to None.
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes. Defaults to None.
            kwargs (dict, optional): Configuration parameters. Defaults to None.
        
        Refer to :py:meth:`~pdf2docx.converter.Converter.default_settings` for detail of 
        configuration parameters.
        
        .. note::
            Change extension from ``pdf`` to ``docx`` if ``docx_file`` is None.
        
        .. note::
            * ``start`` and ``end`` is counted from zero if ``--zero_based_index=True`` (by default).
            * Start from the first page if ``start`` is omitted.
            * End with the last page if ``end`` is omitted.
        
        .. note::
            ``pages`` has a higher priority than ``start`` and ``end``. ``start`` and ``end`` works only
            if ``pages`` is omitted.

        .. note::
            Multi-processing works only for continuous pages specified by ``start`` and ``end`` only.
        """
        settings = self.default_settings
        settings.update(kwargs)
        
        self.parse(start, end, pages, **settings).make_docx(**settings)



    def make_docx(self, **kwargs):
        # check parsed pages
        logging.error(self._color_output('[4/4] Creating pages...'))

        parsed_pages = list(filter(
            lambda page: page.finalized, self._pages
        ))
        if not parsed_pages:
            raise ConversionException(
                'No parsed pages. Please parse page first.'
            )

        # create page by page
        docx_file = Document() 
        for i, page in enumerate(parsed_pages, start=1):        
            if not page.finalized: continue # ignore unparsed pages
            pid = page.id + 1
            try:
                page.make_docx(docx_file)
            except Exception as e:
                logging.error(e)
                if kwargs['ignore_page_error']:
                    logging.error(
                        'Ignore page %d due to making page error: %s', pid, e)
                else:
                    raise MakedocxException(f'Error when make page {pid}: {e}')

        docx_file.save(path_or_stream=self.path_or_stream)


def convert_bytes_to_pdf(stream: bytes, input_type: str):
    """
    You can access files with extensions like .xps, .oxps, .cbz, .fb2 or .epub

    ref: https://pymupdf.readthedocs.io/en/latest/document.html#Document.convert_to_pdf
    """
    # TODO: test 
    with fitz.Document(stream=stream, filetype=input_type) as doc:
        return doc.convert_to_pdf()


perm = int(
	fitz.PDF_PERM_ACCESSIBILITY # always use this
	| fitz.PDF_PERM_PRINT # permit printing
	| fitz.PDF_PERM_COPY # permit copying
	| fitz.PDF_PERM_ANNOTATE # permit annotations
)
encrypt_meth = fitz.PDF_ENCRYPT_AES_256 # strongest algorithm
def encrypt_pdf(file_bytes: bytes, ro_password: str, rw_password: Optional[str] = None):
    # The owner password provides full access rights, including changing passwords, encryption method, or permission detail.
    # The user password provides access to document content according to the established permission details. If present, opening the PDF in a viewer will require providing it.
    # !IMPORTANT: The user password cannot remove pdf password.

    out_bytes = io.BytesIO()
    with fitz.Document(stream=file_bytes, filetype='pdf') as doc:
        if doc.needs_pass:
            raise AlreadyEncryptException(f'This pdf is already encrypted.')

        doc.save(
            out_bytes,
            encryption=encrypt_meth, # set the encryption method
            user_pw=ro_password, # set the user password            
            owner_pw=rw_password, # set the owner password
            permissions=perm, # set permissions
        )
        return out_bytes.getvalue()



def decrypt_pdf(file_bytes: bytes, password:str):
    def is_incorrect_password_and_decrypt():
        # .authenticate() decrypt document
        try:
            return doc.authenticate(password) <= 0
        except Exception as e:
            logging.error(e)
            doc.close()
            return True
    
    out_bytes = io.BytesIO()
    with fitz.Document(stream=file_bytes, filetype='pdf') as doc:
        if not doc.needs_pass:
            return doc.tobytes()

        if is_incorrect_password_and_decrypt():
            raise InvalidPasswordException(f'Invalid password is given {password}')
        
        doc.save(out_bytes)
        return out_bytes.getvalue()




class ConversionException(Exception): 
    pass

class MakedocxException(ConversionException): 
    pass

class InvalidPasswordException(ConversionException):
    pass


class AlreadyEncryptException(ConversionException):
    pass

